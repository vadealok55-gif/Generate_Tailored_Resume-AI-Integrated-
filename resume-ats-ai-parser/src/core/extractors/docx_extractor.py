from docx import Document
from typing import Dict, Any
import logging
from src.core.extractors.base_extractor import BaseExtractor

logger = logging.getLogger(__name__)

class AdvancedDOCXExtractor(BaseExtractor):
    def extract_text(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise
    
    def extract_with_formatting(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            result = {"text": "", "paragraphs": [], "tables": []}
            
            for para in doc.paragraphs:
                result["text"] += para.text + "\n"
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                result["tables"].append(table_data)
            
            return result
        except Exception as e:
            logger.error(f"Error extracting: {e}")
            raise
    
    def extract_structured_data(self, file_path: str) -> Dict[str, Any]:
        return self.extract_with_formatting(file_path)
